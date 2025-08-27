import os
import re
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from data_visualiser import data_visualize
from data_loader import load_current_data
from datetime import datetime

def natural_sort_key(s):
    return [
        int(text) if text.isdigit() else text.lower()
        for text in re.split(r'(\d+)', str(s))
    ]

def generate_report_for_file(file_path: str, doc: Document):
    """
    Генерирует отчёт для одного файла и добавляет его в Word документ.
    
    Аргументы:
    - file_path: путь к CSV файлу с данными
    - doc: экземпляр документа Word
    """
    # Добавляем заголовок с именем файла
    doc.add_heading(f'Отчёт по файлу: {os.path.basename(file_path)}', level=1)
    
    # Загружаем данные
    df = load_current_data(file_path)
    
    # Добавляем описание датафрейма
    doc.add_heading('Статистика данных', level=2)
    
    # Преобразуем описание в читаемый вид
    stats = df.describe().round(3)
    
    # Добавляем таблицу с описательной статистикой
    table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
    table.style = 'Table Grid'
    
    # Заголовки столбцов
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Метрика'
    for i, col in enumerate(stats.columns, 1):
        hdr_cells[i].text = col
    
    # Данные
    for idx, row in stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = idx
        for i, col in enumerate(stats.columns, 1):
            row_cells[i].text = str(row[col])
    
    # Добавляем графики
    doc.add_heading('Визуализация данных', level=2)
    
    # Создаем временный файл для графика
    temp_img = 'temp_plot.png'
    
    try:
        # График токов фаз
        plt.figure(figsize=(10, 6))
        data_visualize(df, 5000) 
        plt.savefig(temp_img, bbox_inches='tight', dpi=100)
        plt.close()
        doc.add_picture(temp_img, width=Inches(6))
        doc.add_paragraph('Рисунок 1: График токов трёх фаз')
        
    finally:
    
        if os.path.exists(temp_img):
            os.remove(temp_img)
    
    doc.add_page_break()

def generate_reports(directory: str, output_file: str = None):
    """
    Генерирует отчёты для всех CSV файлов в указанной директории.
    
    Аргументы:
    - directory: путь к директории с CSV файлами
    - output_file: путь к выходному Word файлу (по умолчанию: reports_<дата_время>.docx)
    """
    if output_file is None:
        timestamp = datetime.now().strftime('%Y%m%d')
        output_file = f'reports_{timestamp}.docx'
    
    doc = Document()
    doc.add_heading('Отчёт по датасету токов электродвигателя', level=0)
    doc.add_paragraph(f'Дата генерации: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    try:
        csv_files = sorted(
            [f for f in os.listdir(directory) if f.endswith('.csv')],
            key=natural_sort_key
        )
        
        if not csv_files:
            print(f'В директории {directory} не найдено CSV файлов')
            return
        
        total_files = len(csv_files)
        for i, filename in enumerate(csv_files, 1):
            file_path = os.path.join(directory, filename)
            print(f'Обработка файла {i}/{total_files}: {filename}')
            try:
                generate_report_for_file(file_path, doc)
            except Exception as e:
                print(f'Ошибка при обработке файла {filename}: {str(e)}')
        
        doc.save(output_file)
        print(f'Отчёт успешно сохранён в файл: {os.path.abspath(output_file)}')
        
    except Exception as e:
        print(f'Ошибка при генерации отчёта: {str(e)}')
        raise

if __name__ == "__main__":
    generate_reports('processed_files')
